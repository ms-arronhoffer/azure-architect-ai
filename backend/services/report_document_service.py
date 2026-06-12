"""Customer-facing report generation — PPTX, DOCX, PDF.

LLM writes the narrative (executive summary, risks, recommendations).
Format-specific builders assemble the final document.
"""

from __future__ import annotations

import json
from datetime import date
from io import BytesIO

from services.openai_service import get_client, get_deployment
from services.pptx_service import build_presentation


# ── LLM narrative ─────────────────────────────────────────────────────────────

_NARRATIVE_TOOL = {
    "type": "function",
    "function": {
        "name": "generate_report_narrative",
        "description": "Write executive narrative sections for an Azure OpenAI model retirement report.",
        "parameters": {
            "type": "object",
            "required": ["executive_summary", "key_risks", "recommended_actions", "customer_highlights"],
            "properties": {
                "executive_summary": {
                    "type": "string",
                    "description": "2-3 sentences summarising the retirement situation and urgency.",
                },
                "key_risks": {
                    "type": "array",
                    "items": {"type": "string"},
                    "description": "4-6 concise bullet points describing key business risks.",
                },
                "recommended_actions": {
                    "type": "array",
                    "items": {"type": "string"},
                    "description": "4-5 prioritised, actionable next steps.",
                },
                "customer_highlights": {
                    "type": "array",
                    "description": "Per-customer narrative for the top priority customers (max 10).",
                    "items": {
                        "type": "object",
                        "required": ["tpid", "name", "summary", "actions"],
                        "properties": {
                            "tpid": {"type": "string"},
                            "name": {"type": "string"},
                            "summary": {"type": "string", "description": "One sentence about this customer's situation."},
                            "actions": {
                                "type": "array",
                                "items": {"type": "string"},
                                "description": "2-3 specific actions for this customer.",
                            },
                        },
                    },
                },
            },
        },
    },
}


def generate_report_narrative(report_data: dict) -> dict:
    """Call LLM once to produce structured executive narrative."""
    summary = report_data.get("summary", {})
    model_summary = report_data.get("model_summary", [])[:12]
    customers = report_data.get("customers", [])[:10]

    compact = {
        "analysis_date": summary.get("analysis_date", str(date.today())),
        "stats": {
            "total_customers": summary.get("unique_customers", 0),
            "total_deployments": summary.get("unique_deployments", 0),
            "unique_models": summary.get("unique_models", 0),
            "critical": summary.get("critical", 0),
            "high": summary.get("high", 0),
            "medium": summary.get("medium", 0),
            "low": summary.get("low", 0),
            "high_usage_urgent": summary.get("high_usage_urgent", 0),
        },
        "models_at_risk": [
            {
                "model": m.get("model", ""),
                "urgency": m.get("urgency", ""),
                "retirement_date": m.get("retirement_date", ""),
                "days_until_retirement": m.get("days_until_retirement"),
                "customers": m.get("customer_count", 0),
                "top_migration": (
                    m["migration_options"][0]["model"]
                    if m.get("migration_options") else ""
                ),
            }
            for m in model_summary
        ],
        "priority_customers": [
            {
                "tpid": c.get("tpid", ""),
                "name": c.get("tp_name", ""),
                "csam": c.get("csam", ""),
                "priority": c.get("priority", ""),
                "deployment_count": len(c.get("deployments", [])),
                "models": list({d["model"] for d in c.get("deployments", [])})[:5],
            }
            for c in customers
        ],
    }

    system = (
        "You are a senior Azure cloud solutions architect writing an executive briefing "
        "for a Microsoft customer-facing model retirement report. Be concise, specific, "
        "and action-oriented. Use business language appropriate for a customer meeting."
    )
    user = (
        "Generate executive narrative for this Azure OpenAI model retirement report:\n\n"
        + json.dumps(compact, indent=2)
    )

    client = get_client()
    deployment = get_deployment("architecture")
    resp = client.chat.completions.create(
        model=deployment,
        messages=[{"role": "system", "content": system}, {"role": "user", "content": user}],
        tools=[_NARRATIVE_TOOL],
        tool_choice={"type": "function", "function": {"name": "generate_report_narrative"}},
        max_completion_tokens=2500,
    )
    tc = resp.choices[0].message.tool_calls[0]
    return json.loads(tc.function.arguments)


# ── PPTX ─────────────────────────────────────────────────────────────────────

def _build_outline(report_data: dict, narrative: dict) -> dict:
    summary = report_data.get("summary", {})
    model_summary = report_data.get("model_summary", [])
    customers = report_data.get("customers", [])
    analysis_date = summary.get("analysis_date", str(date.today()))
    urgent = summary.get("critical", 0) + summary.get("high", 0)

    slides: list[dict] = [
        {
            "layout": "title",
            "title": "Azure OpenAI Model Retirement Analysis",
            "content": [f"Migration Advisor Report  ·  {analysis_date}"],
            "speaker_notes": "Prepared using Azure Migration Advisor.",
        },
        {
            "layout": "quote_stat",
            "title": "Customers Requiring Immediate Action",
            "content": [
                str(urgent),
                (
                    f"{summary.get('unique_models', 0)} models retiring  ·  "
                    f"{summary.get('unique_deployments', 0)} deployments across "
                    f"{summary.get('unique_customers', 0)} customers"
                ),
            ],
        },
        {
            "layout": "content",
            "title": "Executive Summary",
            "content": [narrative.get("executive_summary", "")] + narrative.get("key_risks", [])[:5],
            "speaker_notes": "Key business risks for this retirement cycle.",
        },
        {
            "layout": "two_column",
            "title": "Model Retirement Risk Matrix",
            "content": [
                f"{m['model']}  —  {m.get('urgency', '').upper()}  ({m.get('retirement_date', '?')})"
                for m in model_summary[:8]
            ],
            "right_content": [
                (
                    f"{m['model']}  →  {m['migration_options'][0]['model']} "
                    f"({m['migration_options'][0]['score']}/100)"
                    if m.get("migration_options") else
                    f"{m['model']}  →  No migration data"
                )
                for m in model_summary[:8]
            ],
            "speaker_notes": "Left: models at risk.  Right: top recommended migration target with feasibility score.",
        },
        {
            "layout": "section_divider",
            "title": "Priority Customer Actions",
            "content": [f"Top {min(len(customers), 5)} customers by urgency and usage"],
        },
    ]

    highlights = {h.get("tpid", ""): h for h in narrative.get("customer_highlights", [])}

    for c in customers[:5]:
        tpid = c.get("tpid", "")
        h = highlights.get(tpid, {})
        bullets: list[str] = []
        for d in c.get("deployments", [])[:4]:
            line = f"{d.get('model', '?')}  →  retires {d.get('retirement_date', '?')}"
            opts = d.get("migration_options", [])
            if opts:
                line += f"  |  Migrate to: {opts[0]['model']}"
            bullets.append(line)
        bullets += h.get("actions", [])[:2]
        slides.append({
            "layout": "content",
            "title": f"{c.get('tp_name', 'Unknown')}  ·  TPID {tpid}",
            "content": bullets,
            "speaker_notes": h.get("summary", ""),
        })

    slides.append({
        "layout": "summary",
        "title": "Recommended Next Steps",
        "content": narrative.get("recommended_actions", [])[:5],
    })

    return {"slides": slides}


def build_pptx_report(report_data: dict, narrative: dict) -> bytes:
    return build_presentation(_build_outline(report_data, narrative))


# ── DOCX ─────────────────────────────────────────────────────────────────────

def build_docx_report(report_data: dict, narrative: dict) -> bytes:
    from docx import Document  # type: ignore[import-untyped]
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt, RGBColor as DocxRGB

    AZURE = DocxRGB(0x00, 0x78, 0xD4)
    MUTED = DocxRGB(0x80, 0x98, 0xB0)

    summary = report_data.get("summary", {})
    model_summary = report_data.get("model_summary", [])
    customers = report_data.get("customers", [])
    analysis_date = summary.get("analysis_date", str(date.today()))

    doc = Document()
    sec = doc.sections[0]
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)
    sec.top_margin = Inches(0.9)
    sec.bottom_margin = Inches(0.9)

    # Title
    title_para = doc.add_heading("Azure OpenAI Model Retirement Analysis", 0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in title_para.runs:
        run.font.color.rgb = AZURE
        run.font.size = Pt(24)

    sub = doc.add_paragraph(f"Migration Advisor Report  ·  {analysis_date}")
    sub.runs[0].font.size = Pt(10)
    sub.runs[0].font.color.rgb = MUTED
    doc.add_paragraph()

    # Executive Summary
    doc.add_heading("Executive Summary", 1).runs[0].font.color.rgb = AZURE
    doc.add_paragraph(narrative.get("executive_summary", ""))

    # Urgency stats table
    doc.add_heading("Urgency Overview", 2).runs[0].font.color.rgb = AZURE
    stat_tbl = doc.add_table(rows=2, cols=5)
    stat_tbl.style = "Table Grid"
    for i, hdr in enumerate(["Critical", "High (≤30d)", "Medium (≤90d)", "Low", "High-Usage Urgent"]):
        cell = stat_tbl.cell(0, i)
        cell.text = hdr
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(9)
    for i, val in enumerate([
        summary.get("critical", 0), summary.get("high", 0),
        summary.get("medium", 0), summary.get("low", 0),
        summary.get("high_usage_urgent", 0),
    ]):
        cell = stat_tbl.cell(1, i)
        cell.text = str(val)
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(12)
    doc.add_paragraph()

    # Key risks
    doc.add_heading("Key Risks", 2).runs[0].font.color.rgb = AZURE
    for risk in narrative.get("key_risks", []):
        p = doc.add_paragraph(risk, style="List Bullet")
        if p.runs:
            p.runs[0].font.size = Pt(10)

    doc.add_paragraph()

    # Model Risk Matrix
    doc.add_heading("Model Retirement Risk Matrix", 1).runs[0].font.color.rgb = AZURE
    if model_summary:
        mtbl = doc.add_table(rows=1, cols=5)
        mtbl.style = "Table Grid"
        for i, hdr in enumerate(["Model", "Urgency", "Retires", "Days", "Top Migration"]):
            cell = mtbl.cell(0, i)
            cell.text = hdr
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].runs[0].font.size = Pt(9)
        for m in model_summary:
            row = mtbl.add_row().cells
            days = m.get("days_until_retirement")
            opts = m.get("migration_options", [])
            row[0].text = m.get("model", "")
            row[1].text = m.get("urgency", "").upper()
            row[2].text = m.get("retirement_date", "")
            row[3].text = str(days) if days is not None else "—"
            row[4].text = opts[0]["model"] if opts else "None found"
            for cell in row:
                if cell.paragraphs[0].runs:
                    cell.paragraphs[0].runs[0].font.size = Pt(9)
    doc.add_paragraph()

    # Customer Recommendations
    doc.add_heading("Customer Recommendations", 1).runs[0].font.color.rgb = AZURE
    highlights = {h.get("tpid", ""): h for h in narrative.get("customer_highlights", [])}

    for c in customers:
        tpid = c.get("tpid", "")
        h = highlights.get(tpid, {})

        ch = doc.add_heading(f"{c.get('tp_name', '?')} — TPID {tpid}", 2)
        if ch.runs:
            ch.runs[0].font.size = Pt(12)

        if h.get("summary"):
            p = doc.add_paragraph(h["summary"])
            if p.runs:
                p.runs[0].font.italic = True
                p.runs[0].font.size = Pt(9)

        deps = c.get("deployments", [])
        if deps:
            dtbl = doc.add_table(rows=1, cols=4)
            dtbl.style = "Table Grid"
            for i, hdr in enumerate(["Model", "Retires", "Urgency", "Migrate To"]):
                cell = dtbl.cell(0, i)
                cell.text = hdr
                cell.paragraphs[0].runs[0].font.bold = True
                cell.paragraphs[0].runs[0].font.size = Pt(8)
            for d in deps[:8]:
                row = dtbl.add_row().cells
                opts = d.get("migration_options", [])
                row[0].text = d.get("model", "")
                row[1].text = d.get("retirement_date", "")
                row[2].text = d.get("urgency", "").upper()
                row[3].text = opts[0]["model"] if opts else "—"
                for cell in row:
                    if cell.paragraphs[0].runs:
                        cell.paragraphs[0].runs[0].font.size = Pt(8)

        for action in h.get("actions", []):
            p = doc.add_paragraph(action, style="List Bullet")
            if p.runs:
                p.runs[0].font.size = Pt(9)
        doc.add_paragraph()

    # Next Steps
    doc.add_heading("Recommended Next Steps", 1).runs[0].font.color.rgb = AZURE
    for i, action in enumerate(narrative.get("recommended_actions", []), 1):
        p = doc.add_paragraph(f"{i}. {action}")
        if p.runs:
            p.runs[0].font.size = Pt(10)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── PDF ───────────────────────────────────────────────────────────────────────

def _sanitize_pdf(text: str) -> str:
    """Replace non-Latin-1 characters so fpdf built-in Helvetica doesn't crash."""
    return (
        str(text)
        .replace("\u2014", "-")    # em dash —
        .replace("\u2013", "-")    # en dash –
        .replace("\u00B7", " * ")  # middle dot ·
        .replace("\u2022", "-")    # bullet •
        .replace("\u2019", "'")    # right single quote
        .replace("\u2018", "'")    # left single quote
        .replace("\u201C", '"')    # left double quote
        .replace("\u201D", '"')    # right double quote
        .replace("\u00A0", " ")    # non-breaking space
        .replace("\u2264", "<=")   # ≤
        .replace("\u2265", ">=")   # ≥
        .replace("\u00AE", "(R)")  # ®
        .replace("\u2122", "(TM)") # ™
        .encode("latin-1", errors="replace")
        .decode("latin-1")
    )


def build_pdf_report(report_data: dict, narrative: dict) -> bytes:
    from fpdf import FPDF  # type: ignore[import-untyped]

    summary = report_data.get("summary", {})
    model_summary = report_data.get("model_summary", [])
    customers = report_data.get("customers", [])
    analysis_date = summary.get("analysis_date", str(date.today()))

    AZURE_RGB = (0, 120, 212)
    DARK_RGB = (26, 32, 40)
    CARD_RGB = (35, 43, 53)
    ACCENT_RGB = (80, 194, 255)
    MUTED_RGB = (128, 152, 176)

    URGENCY_RGB: dict[str, tuple] = {
        "CRITICAL": (200, 50, 50),
        "HIGH": (220, 100, 0),
        "MEDIUM": (180, 140, 0),
        "LOW": (0, 160, 80),
        "UNKNOWN": (100, 100, 100),
    }

    class _PDF(FPDF):
        def header(self):
            self.set_fill_color(*AZURE_RGB)
            self.rect(0, 0, 210, 5, "F")
            self.set_font("Helvetica", size=7)
            self.set_text_color(255, 255, 255)
            self.set_xy(10, 0)
            self.cell(190, 5, "Azure OpenAI Model Retirement Analysis - Confidential", align="L")
            self.set_text_color(0, 0, 0)
            self.ln(7)

        def footer(self):
            self.set_y(-11)
            self.set_font("Helvetica", size=7)
            self.set_text_color(*MUTED_RGB)
            self.cell(0, 5, f"Generated {analysis_date} | Azure Migration Advisor | Page {self.page_no()}", align="C")
            self.set_text_color(0, 0, 0)

    pdf = _PDF(orientation="P", unit="mm", format="A4")
    pdf.set_margins(15, 16, 15)
    pdf.set_auto_page_break(auto=True, margin=14)
    pdf.add_page()

    # ── Title block
    pdf.set_fill_color(*DARK_RGB)
    pdf.rect(15, 17, 180, 26, "F")
    pdf.set_fill_color(*AZURE_RGB)
    pdf.rect(15, 17, 3, 26, "F")
    pdf.set_font("Helvetica", "B", 18)
    pdf.set_text_color(255, 255, 255)
    pdf.set_xy(21, 20)
    pdf.cell(170, 9, "Azure OpenAI Model Retirement Analysis", ln=True)
    pdf.set_font("Helvetica", "", 9)
    pdf.set_text_color(*ACCENT_RGB)
    pdf.set_x(21)
    pdf.cell(170, 7, f"Migration Advisor Report | {analysis_date}", ln=True)
    pdf.ln(10)
    pdf.set_text_color(0, 0, 0)

    # ── Executive Summary
    pdf.set_font("Helvetica", "B", 12)
    pdf.set_text_color(*AZURE_RGB)
    pdf.cell(0, 6, "Executive Summary", ln=True)
    pdf.set_draw_color(*AZURE_RGB)
    pdf.line(15, pdf.get_y(), 195, pdf.get_y())
    pdf.ln(2)
    pdf.set_font("Helvetica", "", 9)
    pdf.set_text_color(0, 0, 0)
    pdf.multi_cell(0, 5, _sanitize_pdf(narrative.get("executive_summary", "")))
    pdf.ln(4)

    # ── Stat cards
    stats = [
        ("Critical", summary.get("critical", 0), URGENCY_RGB["CRITICAL"]),
        ("High (<=30d)", summary.get("high", 0), URGENCY_RGB["HIGH"]),
        ("Medium (<=90d)", summary.get("medium", 0), URGENCY_RGB["MEDIUM"]),
        ("Low", summary.get("low", 0), URGENCY_RGB["LOW"]),
        ("High-Usage Urgent", summary.get("high_usage_urgent", 0), URGENCY_RGB["CRITICAL"]),
    ]
    card_w = 35
    y0 = pdf.get_y()
    for idx, (label, val, color) in enumerate(stats):
        x = 15 + idx * (card_w + 1)
        pdf.set_fill_color(*CARD_RGB)
        pdf.rect(x, y0, card_w, 16, "F")
        pdf.set_fill_color(*color)
        pdf.rect(x, y0, 2, 16, "F")
        pdf.set_font("Helvetica", "B", 15)
        pdf.set_text_color(*color)
        pdf.set_xy(x + 4, y0 + 1)
        pdf.cell(card_w - 6, 8, str(val))
        pdf.set_font("Helvetica", "", 6)
        pdf.set_text_color(*MUTED_RGB)
        pdf.set_xy(x + 4, y0 + 9)
        pdf.cell(card_w - 6, 5, label)
    pdf.set_y(y0 + 19)
    pdf.set_text_color(0, 0, 0)

    # ── Key Risks
    pdf.ln(2)
    pdf.set_font("Helvetica", "B", 10)
    pdf.set_text_color(*AZURE_RGB)
    pdf.cell(0, 6, "Key Risks", ln=True)
    pdf.set_text_color(0, 0, 0)
    for risk in narrative.get("key_risks", []):
        pdf.set_font("Helvetica", "", 9)
        pdf.set_x(17)
        pdf.cell(4, 5, "-")
        pdf.multi_cell(0, 5, _sanitize_pdf(risk))
    pdf.ln(4)

    # ── Model Risk Matrix
    pdf.set_font("Helvetica", "B", 12)
    pdf.set_text_color(*AZURE_RGB)
    pdf.cell(0, 6, "Model Retirement Risk Matrix", ln=True)
    pdf.set_draw_color(*AZURE_RGB)
    pdf.line(15, pdf.get_y(), 195, pdf.get_y())
    pdf.ln(2)
    pdf.set_text_color(0, 0, 0)

    if model_summary:
        col_w = [52, 22, 28, 16, 62]
        hdrs  = ["Model", "Urgency", "Retires", "Days", "Top Migration"]
        pdf.set_fill_color(*DARK_RGB)
        pdf.set_text_color(255, 255, 255)
        pdf.set_font("Helvetica", "B", 7)
        for w, h in zip(col_w, hdrs):
            pdf.cell(w, 5, h, border=0, fill=True)
        pdf.ln()
        pdf.set_text_color(0, 0, 0)
        for idx, m in enumerate(model_summary):
            fill = (240, 244, 248) if idx % 2 == 0 else (255, 255, 255)
            pdf.set_fill_color(*fill)
            pdf.set_font("Helvetica", "", 7)
            opts = m.get("migration_options", [])
            days = m.get("days_until_retirement")
            urgency = m.get("urgency", "").upper()
            row = [
                _sanitize_pdf(m.get("model", ""))[:30],
                _sanitize_pdf(urgency),
                _sanitize_pdf(m.get("retirement_date", "")),
                str(days) if days is not None else "-",
                _sanitize_pdf(opts[0]["model"] if opts else "None found")[:35],
            ]
            for w, val in zip(col_w, row):
                pdf.cell(w, 5, val, border=0, fill=True)
            pdf.ln()
    pdf.ln(5)

    # ── Customer Recommendations
    pdf.add_page()
    pdf.set_font("Helvetica", "B", 12)
    pdf.set_text_color(*AZURE_RGB)
    pdf.cell(0, 6, "Priority Customer Recommendations", ln=True)
    pdf.set_draw_color(*AZURE_RGB)
    pdf.line(15, pdf.get_y(), 195, pdf.get_y())
    pdf.ln(3)
    pdf.set_text_color(0, 0, 0)

    highlights = {h.get("tpid", ""): h for h in narrative.get("customer_highlights", [])}

    for c in customers[:10]:
        tpid = c.get("tpid", "")
        h = highlights.get(tpid, {})
        priority = c.get("priority", "unknown").upper()
        bar_color = URGENCY_RGB.get(priority, URGENCY_RGB["UNKNOWN"])

        y0 = pdf.get_y()
        pdf.set_fill_color(*bar_color)
        pdf.rect(15, y0, 3, 7, "F")
        pdf.set_fill_color(*CARD_RGB)
        pdf.rect(18, y0, 177, 7, "F")
        pdf.set_font("Helvetica", "B", 9)
        pdf.set_text_color(255, 255, 255)
        pdf.set_xy(21, y0)
        pdf.cell(140, 7, _sanitize_pdf(f"{c.get('tp_name', '?')}  -  TPID {tpid}"))
        pdf.set_font("Helvetica", "", 7)
        pdf.set_text_color(*bar_color)
        pdf.cell(32, 7, priority, align="R", ln=True)
        pdf.set_text_color(0, 0, 0)
        pdf.ln(1)

        if h.get("summary"):
            pdf.set_font("Helvetica", "I", 8)
            pdf.set_text_color(80, 80, 80)
            pdf.set_x(18)
            pdf.multi_cell(0, 4, _sanitize_pdf(h["summary"]))
            pdf.set_text_color(0, 0, 0)

        for d in c.get("deployments", [])[:5]:
            opts = d.get("migration_options", [])
            migrate_to = _sanitize_pdf(opts[0]["model"] if opts else "-")
            days = d.get("days_until_retirement")
            days_str = (
                f"{days}d" if days is not None and days >= 0
                else ("PAST" if days is not None else "—")
            )
            pdf.set_x(20)
            pdf.set_fill_color(*AZURE_RGB)
            pdf.rect(20, pdf.get_y() + 1, 2, 3, "F")
            pdf.set_font("Helvetica", "", 7)
            pdf.cell(78, 5, _sanitize_pdf(d.get("model", ""))[:42])
            pdf.set_text_color(100, 100, 100)
            pdf.cell(28, 5, _sanitize_pdf(d.get("retirement_date", "")))
            pdf.cell(18, 5, days_str)
            pdf.set_text_color(*AZURE_RGB)
            pdf.cell(0, 5, migrate_to[:32], ln=True)
            pdf.set_text_color(0, 0, 0)

        for action in h.get("actions", []):
            pdf.set_font("Helvetica", "", 8)
            pdf.set_x(20)
            pdf.cell(4, 4, "-")
            pdf.multi_cell(0, 4, _sanitize_pdf(action))
        pdf.ln(3)

    # ── Next Steps
    if pdf.get_y() > 245:
        pdf.add_page()
    pdf.set_font("Helvetica", "B", 12)
    pdf.set_text_color(*AZURE_RGB)
    pdf.cell(0, 6, "Recommended Next Steps", ln=True)
    pdf.set_draw_color(*AZURE_RGB)
    pdf.line(15, pdf.get_y(), 195, pdf.get_y())
    pdf.ln(3)
    pdf.set_text_color(0, 0, 0)

    for i, action in enumerate(narrative.get("recommended_actions", []), 1):
        y0 = pdf.get_y()
        pdf.set_fill_color(*AZURE_RGB)
        pdf.rect(15, y0, 7, 6, "F")
        pdf.set_text_color(255, 255, 255)
        pdf.set_font("Helvetica", "B", 8)
        pdf.set_xy(15, y0)
        pdf.cell(7, 6, str(i), align="C")
        pdf.set_text_color(0, 0, 0)
        pdf.set_font("Helvetica", "", 9)
        pdf.set_x(24)
        pdf.multi_cell(0, 6, _sanitize_pdf(action))
        pdf.ln(1)

    return bytes(pdf.output())
